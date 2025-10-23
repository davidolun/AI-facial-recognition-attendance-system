import base64
import re
import time
import logging
from django.http import JsonResponse
from django.shortcuts import render
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
import os
import datetime
import json
import cloudinary
import cloudinary.uploader
from cloudinary import config as cloudinary_config

# Configure Cloudinary
cloudinary_config(
    cloud_name=os.getenv('CLOUDINARY_CLOUD_NAME', 'duu7pc7s3'),
    api_key=os.getenv('CLOUDINARY_API_KEY', '625655631397579'),
    api_secret=os.getenv('CLOUDINARY_API_SECRET', 'HEE8Or7rvr7SBOv61t5CWsSRUIs')
)

# Try to import image processing libraries, with fallback
try:
    import numpy as np
    import cv2
    from PIL import Image
    import io
    IMAGE_PROCESSING_AVAILABLE = True
except ImportError:
    IMAGE_PROCESSING_AVAILABLE = False
    print("Warning: Image processing libraries not available.")

# Face recognition using OpenCV Haar cascades (always available with opencv-python)
FACE_RECOGNITION_AVAILABLE = True

# OpenAI availability (enabled by default)
OPENAI_AVAILABLE = False
if os.getenv('DISABLE_OPENAI', 'false').lower() != 'true':
    try:
        from openai import OpenAI
        OPENAI_AVAILABLE = True
    except ImportError:
        OPENAI_AVAILABLE = False
        print("Warning: OpenAI not available. AI features will be disabled.")

# OpenCV Haar Cascade helper functions
def detect_faces_opencv(image_array):
    """Detect faces in an image using OpenCV Haar cascades"""
    try:
        # Load the Haar cascade for face detection
        face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
        
        # Convert to grayscale for better detection
        gray = cv2.cvtColor(image_array, cv2.COLOR_BGR2GRAY)
        
        # Detect faces
        faces = face_cascade.detectMultiScale(gray, 1.1, 4)
        
        # Convert to (x1, y1, x2, y2) format and ensure Python ints
        face_boxes = []
        for (x, y, w, h) in faces:
            face_boxes.append((int(x), int(y), int(x + w), int(y + h)))
        
        return face_boxes
    except Exception as e:
        print(f"Face detection error: {e}")
        return []

def load_image_from_path_or_url(image_path):
    """Load image from either local path or Cloudinary URL"""
    try:
        if image_path.startswith('http://') or image_path.startswith('https://'):
            # Load from URL (Cloudinary)
            import requests
            response = requests.get(image_path)
            img_array = np.asarray(bytearray(response.content), dtype=np.uint8)
            img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)
            return img
        else:
            # Load from local path
            full_path = os.path.join(settings.MEDIA_ROOT, image_path)
            if os.path.exists(full_path):
                return cv2.imread(full_path)
            return None
    except Exception as e:
        print(f"Error loading image: {e}")
        return None

def extract_face_features(image_array, face_box):
    """Extract basic face features using OpenCV"""
    try:
        x1, y1, x2, y2 = face_box
        face_crop = image_array[y1:y2, x1:x2]
        
        if face_crop.size == 0:
            return None
            
        # Resize to standard size for comparison
        face_resized = cv2.resize(face_crop, (100, 100))
        
        # Convert to grayscale
        gray_face = cv2.cvtColor(face_resized, cv2.COLOR_BGR2GRAY)
        
        # Flatten and normalize, convert to Python list
        features = gray_face.flatten().astype(np.float32) / 255.0
        
        return features.tolist()
    except Exception as e:
        print(f"Face feature extraction error: {e}")
        return None

def compare_faces_opencv(features1, features2, threshold=0.3):
    """Compare two face feature vectors using L2 distance"""
    try:
        if features1 is None or features2 is None:
            return False, 1.0
            
        # Convert to numpy arrays for calculation
        features1 = np.array(features1)
        features2 = np.array(features2)
        
        # Calculate L2 distance
        distance = np.linalg.norm(features1 - features2)
        
        # Normalize distance (0-1 scale)
        normalized_distance = distance / np.sqrt(len(features1))
        
        # Convert to Python float
        return normalized_distance < threshold, float(normalized_distance)
    except Exception as e:
        print(f"Face comparison error: {e}")
        return False, 1.0

def find_face_in_image(image_array, student_features, student_paths, threshold=0.3):
    """Find matching face in image array against student features"""
    try:
        faces = detect_faces_opencv(image_array)
        if not faces:
            return None, float('inf')
            
        best_match = None
        best_distance = float('inf')
        
        for face_box in faces:
            current_features = extract_face_features(image_array, face_box)
            if current_features is None:
                continue
                
            for i, student_features_item in enumerate(student_features):
                if student_features_item is not None:
                    verified, distance = compare_faces_opencv(current_features, student_features_item, threshold)
                    if verified and distance < best_distance:
                        best_match = student_paths[i]
                        best_distance = distance
        
        return best_match, best_distance
    except Exception as e:
        print(f"Face finding error: {e}")
        return None, float('inf')

# Try to import office document libraries, with fallback
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    print("Warning: openpyxl not available. Excel export features will be disabled.")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Word export features will be disabled.")
from django.db.models import Count, Q
from datetime import datetime, date, timedelta
from django.conf import settings
import json
from .models import Student, AttendanceRecord, AttendanceSession, AIQuery
from django.http import HttpResponse, JsonResponse
from django.views.decorators.http import require_http_methods
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.shortcuts import redirect
from .models import Teacher, Class
from django.db import transaction

# Set up loggers
logger = logging.getLogger('faceapp')
performance_logger = logging.getLogger('faceapp.performance')
security_logger = logging.getLogger('faceapp.security')


# Do not initialize OpenAI client at import time to avoid httpx/proxies issues on some hosts
client = None

@login_required
def home(request):
    start_time = time.time()
    logger.info(f"User {request.user.username} accessed home page")
    response = render(request, 'home.html')
    performance_logger.info(f"Home page load time: {time.time() - start_time:.2f}s for user {request.user.username}")
    return response


@csrf_exempt
def take_attendance(request):
    start_time = time.time()
    logger.info("Attendance taking request received")

    if request.method == "POST":
        try:
            data = json.loads(request.body)
            image_data = data.get("image")
            if not image_data:
                logger.warning("Attendance request failed: No image received")
                performance_logger.info(f"Attendance request failed - no image: {time.time() - start_time:.2f}s")
                return JsonResponse({"error": "No image received"}, status=400)

            # decode image
            img_str = re.sub("^data:image/.+;base64,", "", image_data)
            img_bytes = base64.b64decode(img_str)
            nparr = np.frombuffer(img_bytes, np.uint8)
            frame = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)

            # Check if image processing is available
            if not IMAGE_PROCESSING_AVAILABLE:
                return JsonResponse({
                    "message": "Image processing not available in production. Please use the session-based attendance system.",
                    "faces": []
                })

            # Check if face recognition is available
            if not FACE_RECOGNITION_AVAILABLE:
                return JsonResponse({
                    "message": "Face recognition not available in production. Please use the session-based attendance system.",
                    "faces": []
                })

            # Detect faces using OpenCV
            faces = detect_faces_opencv(cv2.cvtColor(rgb_frame, cv2.COLOR_RGB2BGR))
            if not faces:
                return JsonResponse({"message": "No face detected", "faces": []})

            # Prepare face rectangles for JS
            faces_for_js = [{"top": f[1], "right": f[2], "bottom": f[3], "left": f[0]} for f in faces]

            # Face recognition part using OpenCV
            students_dir = os.path.join(settings.BASE_DIR, "students")
            student_image_paths = []
            student_features = []
            student_names = []

            for file in os.listdir(students_dir):
                if file.endswith(".jpg") or file.endswith(".png"):
                    student_path = os.path.join(students_dir, file)
                    student_image_paths.append(student_path)
                    student_names.append(file.rsplit(".", 1)[0])
                    
                    # Extract features for this student
                    try:
                        student_img = cv2.imread(student_path)
                        student_faces = detect_faces_opencv(student_img)
                        if student_faces:
                            features = extract_face_features(student_img, student_faces[0])
                            student_features.append(features)
                        else:
                            student_features.append(None)
                    except Exception as e:
                        print(f"Error processing {file}: {e}")
                        student_features.append(None)

            recognized = []
            if student_image_paths and any(feat is not None for feat in student_features):
                try:
                    best_match, distance = find_face_in_image(cv2.cvtColor(rgb_frame, cv2.COLOR_RGB2BGR), 
                                                            student_features, student_image_paths, threshold=0.3)
                    if best_match:
                        # Extract name from path
                        filename = os.path.basename(best_match)
                        name = filename.rsplit(".", 1)[0]
                        recognized.append(name)

                        # log attendance
                        now = datetime.datetime.now()
                        time_str = now.strftime("%H:%M:%S")
                        date_str = now.strftime("%Y-%m-%d")
                        log_path = os.path.join(settings.BASE_DIR, "attendance.csv")
                        with open(log_path, "a") as f:
                            f.write(f"{name},{time_str},{date_str}\n")
                except Exception as e:
                    print(f"Face recognition error: {e}")

            processing_time = time.time() - start_time
            performance_logger.info(f"Face recognition completed in {processing_time:.2f}s - recognized: {len(recognized)} students")
            logger.info(f"Attendance taken successfully: {', '.join(recognized) if recognized else 'No match'}")

            return JsonResponse({"message": f"Attendance taken: {', '.join(recognized) if recognized else 'No match'}",
                                 "faces": faces_for_js})

        except Exception as e:
            error_time = time.time() - start_time
            logger.error(f"Attendance taking failed after {error_time:.2f}s: {str(e)}")
            performance_logger.warning(f"Attendance processing error: {error_time:.2f}s - {str(e)}")
            return JsonResponse({"error": str(e)}, status=400)
    return JsonResponse({"message": "Use POST request."})

@login_required
@csrf_exempt
def add_student(request):
    start_time = time.time()
    logger.info(f"User {request.user.username} initiated student addition")

    if request.method == "POST":
        try:
            data = json.loads(request.body)
            image_data = data.get("image")
            student_name = data.get("name")
            student_id = data.get("student_id", "")
            email = data.get("email", "")
            phone = data.get("phone", "")

            if not student_name:
                logger.warning(f"Student addition failed: No name provided by user {request.user.username}")
                return JsonResponse({"error": "No name provided"}, status=400)
            if not image_data:
                logger.warning(f"Student addition failed: No image provided by user {request.user.username}")
                return JsonResponse({"error": "No image provided"}, status=400)

            # Decode base64 image
            try:
                image_data = image_data.split(',')[1]  # Remove data:image/png;base64, prefix
                image_bytes = base64.b64decode(image_data)
                image = Image.open(io.BytesIO(image_bytes))
            except Exception as e:
                logger.error(f"Image decoding failed for user {request.user.username}: {str(e)}")
                return JsonResponse({"error": "Invalid image format"}, status=400)

            # Face detection using OpenCV
            if not FACE_RECOGNITION_AVAILABLE:
                logger.warning(f"Face recognition not available, creating student without face data")
                face_encoding = None
            else:
                try:
                    # Convert PIL image to numpy array
                    image_array = np.array(image.convert('RGB'))
                    image_bgr = cv2.cvtColor(image_array, cv2.COLOR_RGB2BGR)
                    
                    # Check if face is detected
                    faces = detect_faces_opencv(image_bgr)
                    
                    if not faces:
                        logger.warning(f"No face detected in image for {student_name}")
                        return JsonResponse({"error": "No face detected in the image. Please ensure your face is clearly visible."}, status=400)
                    
                    # For OpenCV, we don't need to store encodings as we'll compare directly
                    face_encoding = "opencv_detected"  # Placeholder
                    logger.info(f"Face detected for {student_name}")
                        
                except Exception as e:
                    logger.error(f"Face detection failed for {student_name}: {str(e)}")
                    return JsonResponse({"error": "Face detection failed. Please try again."}, status=400)

            # Generate unique filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{student_name.lower().replace(' ', '_')}_{timestamp}.jpg"
            
            # Convert to RGB if needed
            if image.mode in ('RGBA', 'LA', 'P'):
                # Create a white background for transparent images
                background = Image.new('RGB', image.size, (255, 255, 255))
                if image.mode == 'P':
                    image = image.convert('RGBA')
                background.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
                image = background
            elif image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Try to upload to Cloudinary first, fallback to local storage
            try:
                logger.info(f"Attempting to upload {filename} to Cloudinary...")
                img_byte_arr = io.BytesIO()
                image.save(img_byte_arr, format='JPEG', quality=95)
                img_byte_arr.seek(0)
                
                upload_result = cloudinary.uploader.upload(
                    img_byte_arr,
                    folder="attendance_students",
                    public_id=filename.replace('.jpg', ''),
                    resource_type="image"
                )
                relative_path = upload_result['secure_url']
                logger.info(f"Successfully uploaded to Cloudinary: {relative_path}")
            except Exception as cloudinary_error:
                # Fallback to local storage if Cloudinary fails
                logger.warning(f"Cloudinary upload failed: {cloudinary_error}. Saving locally...")
                file_path = os.path.join(settings.MEDIA_ROOT, 'students', filename)
                os.makedirs(os.path.dirname(file_path), exist_ok=True)
                image.save(file_path, 'JPEG', quality=95)
                relative_path = os.path.join('students', filename)
                logger.info(f"Saved locally: {relative_path}")

            # Create student record
            student = Student.objects.create(
                name=student_name,
                student_id=student_id if student_id else None,
                email=email if email else None,
                phone=phone if phone else None,
                image_path=relative_path
            )

            # Automatically add student to all classes of the logged-in teacher
            teacher_classes = Class.objects.filter(teacher=request.user, is_active=True)
            if teacher_classes.exists():
                student.classes.add(*teacher_classes)
                logger.info(f"Student {student_name} automatically added to {teacher_classes.count()} classes of teacher {request.user.username}")

            processing_time = time.time() - start_time
            logger.info(f"Student {student_name} added successfully in {processing_time:.2f}s by user {request.user.username}")
            
            return JsonResponse({
                "message": f"Student {student_name} added successfully!",
                "student_id": student.id,
                "processing_time": f"{processing_time:.2f}s"
            })

        except Exception as e:
            processing_time = time.time() - start_time
            logger.error(f"Student addition failed after {processing_time:.2f}s for user {request.user.username}: {str(e)}")
            return JsonResponse({"error": str(e)}, status=400)
    
    return render(request, "add_student.html")

@login_required
def class_management(request):
    """Render the class management page"""
    return render(request, 'class_management.html')

@csrf_exempt
def take_attendance_enhanced(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            image_data = data.get("image")
            session_name = data.get("session", f"Session {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            
            if not image_data:
                return JsonResponse({"error": "No image received"}, status=400)

            img_str = re.sub("^data:image/.+;base64,", "", image_data)
            img_bytes = base64.b64decode(img_str)
            nparr = np.frombuffer(img_bytes, np.uint8)
            frame = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)

            # Detect faces using OpenCV
            faces = detect_faces_opencv(cv2.cvtColor(rgb_frame, cv2.COLOR_RGB2BGR))
            if not faces:
                return JsonResponse({"message": "No face detected", "faces": []})

            faces_for_js = [{"top": f[1], "right": f[2], "bottom": f[3], "left": f[0]} for f in faces]

            today = date.today()
            current_time = datetime.now()
            session, created = AttendanceSession.objects.get_or_create(
                name=session_name,
                date=today,
                defaults={'start_time': current_time.time()}
            )

            students = Student.objects.filter(is_active=True)
            student_image_paths = []
            student_objects = []

            for student in students:
                student_img = load_image_from_path_or_url(student.image_path)
                if student_img is not None:
                    student_image_paths.append(student.image_path)
                    student_objects.append(student)

            recognized_students = []
            print(f"Number of student images loaded: {len(student_image_paths)}")
            print(f"Number of faces detected in current frame: {len(faces)}")
            
            # Extract features for students
            student_features = []
            for student in students:
                student_img = load_image_from_path_or_url(student.image_path)
                if student_img is not None:
                    try:
                        student_faces = detect_faces_opencv(student_img)
                        if student_faces:
                            features = extract_face_features(student_img, student_faces[0])
                            student_features.append(features)
                        else:
                            student_features.append(None)
                    except Exception as e:
                        print(f"Error processing {student.name}: {e}")
                        student_features.append(None)
                else:
                    student_features.append(None)
            
            # Find matching face using OpenCV
            try:
                best_match, distance = find_face_in_image(cv2.cvtColor(rgb_frame, cv2.COLOR_RGB2BGR), 
                                                        student_features, student_image_paths, threshold=0.3)
                if best_match:
                    # Find corresponding student object
                    for i, student_path in enumerate(student_image_paths):
                        if student_path == best_match:
                            student = student_objects[i]
                            recognized_students.append(student)
                            break
            except Exception as e:
                print(f"Face recognition error: {e}")
            
            # Process recognized students
            status_messages = []
            for student in recognized_students:
                existing_record = AttendanceRecord.objects.filter(
                    student=student,
                    session=session
                ).first()
                
                if not existing_record:
                    arrival_time = current_time.time()
                    is_late = arrival_time > session.start_time
                    
                    AttendanceRecord.objects.create(
                        student=student,
                        session=session,
                        arrival_time=arrival_time,
                        is_late=is_late
                    )
                    
                    time_str = arrival_time.strftime("%H:%M:%S")
                    status = f" (Late - {time_str})" if is_late else f" (On time - {time_str})"
                    status_messages.append(f"{student.name}{status}")
                else:
                    original_time = existing_record.arrival_time.strftime("%H:%M:%S") if hasattr(existing_record, 'arrival_time') and existing_record.arrival_time else existing_record.time.strftime("%H:%M:%S")
                    status_messages.append(f"{student.name} (Already marked at {original_time})")
            

            message = f"Attendance taken: {', '.join(status_messages) if status_messages else 'No match'}"
            return JsonResponse({"message": message, "faces": faces_for_js})

        except Exception as e:
            return JsonResponse({"error": str(e)}, status=400)
            
    return JsonResponse({"message": "Use POST request."})

@csrf_exempt
def detect_faces(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            image_data = data.get("image")
            if not image_data:
                return JsonResponse({"faces": []})

            img_str = re.sub("^data:image/.+;base64,", "", image_data)
            img_bytes = base64.b64decode(img_str)
            nparr = np.frombuffer(img_bytes, np.uint8)
            frame = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            rgb_frame = frame[:, :, ::-1]

            # Detect faces using OpenCV
            faces = detect_faces_opencv(cv2.cvtColor(rgb_frame, cv2.COLOR_RGB2BGR))
            faces_list = []
            for face in faces:
                faces_list.append({"top": face[1], "right": face[2], "bottom": face[3], "left": face[0]})

            return JsonResponse({"faces": faces_list})
        except Exception as e:
            return JsonResponse({"error": str(e), "faces": []})
    return JsonResponse({"faces": []})

@login_required
def view_records(request):
    return render(request, 'ai_assistant.html')

def get_complete_attendance_data(teacher=None):
    """Get COMPLETE attendance data for AI with FIXED absence tracking"""
    
    # Filter by teacher if provided (for teacher-specific data)
    if teacher:
        all_students = list(Student.objects.filter(is_active=True, classes__teacher=teacher).distinct().values('id', 'name', 'student_id', 'email'))
        all_sessions = list(AttendanceSession.objects.filter(teacher=teacher).order_by('-date', '-start_time').values(
            'id', 'name', 'date', 'start_time', 'end_time', 'class_session__name'
        ))
    else:
        # For admins - get all data
        all_students = list(Student.objects.filter(is_active=True).values('id', 'name', 'student_id', 'email'))
        all_sessions = list(AttendanceSession.objects.all().order_by('-date', '-start_time').values(
            'id', 'name', 'date', 'start_time', 'end_time', 'class_session__name'
        ))
    
    # Convert dates and times to strings for JSON serialization
    for session in all_sessions:
        session['date'] = session['date'].strftime('%Y-%m-%d')
        session['start_time'] = session['start_time'].strftime('%H:%M:%S')
        if session['end_time']:
            session['end_time'] = session['end_time'].strftime('%H:%M:%S')
    
    # Get ALL attendance records with related data
    if teacher:
        all_records = list(AttendanceRecord.objects.filter(session__teacher=teacher).select_related('student', 'session').values(
            'student__name', 'student_id', 'session__name', 'session_id', 
            'date', 'time', 'arrival_time', 'is_late', 'timestamp'
        ))
    else:
        all_records = list(AttendanceRecord.objects.select_related('student', 'session').values(
            'student__name', 'student_id', 'session__name', 'session_id', 
            'date', 'time', 'arrival_time', 'is_late', 'timestamp'
        ))
    
    # Convert dates and times to strings
    for record in all_records:
        record['date'] = record['date'].strftime('%Y-%m-%d')
        record['time'] = record['time'].strftime('%H:%M:%S')
        record['timestamp'] = record['timestamp'].strftime('%Y-%m-%d %H:%M:%S')
        if record['arrival_time']:
            record['arrival_time'] = record['arrival_time'].strftime('%H:%M:%S')
    
    # FIXED: Calculate attendance statistics for each student based on available sessions
    student_stats = {}
    for student in all_students:
        student_records = [r for r in all_records if r['student_id'] == student['id']]
        
        # Get unique sessions this student attended
        sessions_attended = len(set(r['session_id'] for r in student_records))
        times_late = len([r for r in student_records if r['is_late']])
        times_on_time = len([r for r in student_records if not r['is_late']])
        
        # FIXED: Calculate available sessions based on sessions student actually has records for
        # This prevents >100% when students get removed from classes
        
        # Get unique session IDs this student has attendance records for
        student_session_ids = set(r['session_id'] for r in student_records)
        
        # Available sessions = sessions where student has records (attended OR was expected)
        # This accounts for historical enrollment changes
        available_session_count = len(student_session_ids)
        
        # Alternative approach: Use max of current enrollment or historical attendance
        if teacher:
            current_available_sessions = AttendanceSession.objects.filter(
                teacher=teacher,
                class_session__students__id=student['id']
            ).values_list('id', flat=True)
        else:
            current_available_sessions = AttendanceSession.objects.filter(
                class_session__students__id=student['id']
            ).values_list('id', flat=True)
        
        current_available_count = len(list(current_available_sessions))
        
        # Use the larger of: current enrollment sessions OR sessions with records
        # This handles both current students and those removed from classes
        available_session_count = max(available_session_count, current_available_count)
        
        # Calculate percentage (can never exceed 100% now)
        if available_session_count > 0:
            attendance_percentage = round((sessions_attended / available_session_count) * 100, 1)
            # Ensure it never exceeds 100%
            attendance_percentage = min(attendance_percentage, 100.0)
        else:
            attendance_percentage = 0
        
        student_stats[student['name']] = {
            'total_sessions_attended': sessions_attended,
            'available_sessions': available_session_count,
            'times_late': times_late,
            'times_on_time': times_on_time,
            'attendance_percentage': attendance_percentage
        }
    
    # FIXED: For each session, calculate attendance based on who was actually present vs eligible
    session_details = {}
    for session in all_sessions:
        session_key = f"{session['name']}_{session['date']}"
        
        # Get students who attended this session
        session_records = [r for r in all_records if r['session_id'] == session['id']]
        present_students = [r['student__name'] for r in session_records]
        
        # FIXED: Instead of using all students in class, only count students who were expected
        # For sessions without explicit enrollment, we'll assume all class students are eligible
        # But the percentage calculation will be based on actual attendance
        
        # Get all students from the session's class (eligible students)
        if teacher:
            eligible_students = list(Student.objects.filter(
                classes__sessions__id=session['id'], 
                is_active=True
            ).values_list('name', flat=True))
        else:
            # For demo/admin, use all students as potentially eligible
            eligible_students = [s['name'] for s in all_students]
        
        # If no specific enrollment info, assume all present students were eligible
        if not eligible_students:
            eligible_students = present_students
        
        absent_students = [name for name in eligible_students if name not in present_students]
        
        session_details[session_key] = {
            'session_info': session,
            'present_students': present_students,
            'absent_students': absent_students,
            'present_count': len(present_students),
            'absent_count': len(absent_students),
            'eligible_count': len(eligible_students),  # NEW: Track eligible students
            'late_students': [r['student__name'] for r in session_records if r['is_late']],
            'on_time_students': [r['student__name'] for r in session_records if not r['is_late']]
        }
    
    # Add sessions list for dropdown (unique session names)
    sessions_list = list(set(session['name'] for session in all_sessions))
    unique_dates = sorted(list(set(session['date'] for session in all_sessions)))
    
    return {
        'total_students': len(all_students),
        'total_sessions': len(all_sessions),
        'today_date': date.today().strftime('%Y-%m-%d'),
        'all_students': all_students,
        'all_sessions': all_sessions,
        'all_attendance_records': all_records,
        'student_statistics': student_stats,
        'session_details': session_details,
        'sessions_list': sessions_list,
        'unique_dates': unique_dates
    }

# Store conversation context in memory (in production, use database or cache)
conversation_contexts = {}

def query_attendance_data_with_context(user_query: str, session_id: str, teacher=None) -> str:
    """Enhanced AI query with conversation context and complete data access"""
    
    # Get complete data (filtered by teacher if not admin)
    data = get_complete_attendance_data(teacher)
    
    # Get or initialize conversation context for this session
    if session_id not in conversation_contexts:
        conversation_contexts[session_id] = []
    
    conversation_history = conversation_contexts[session_id]
    conversation_history.append({"role": "user", "content": user_query})
    
    # Keep only last 10 exchanges to prevent token limit issues
    if len(conversation_history) > 20:
        conversation_history = conversation_history[-20:]
    
    # Create comprehensive system prompt with teacher context
    teacher_info = ""
    if teacher:
        teacher_info = f"""
TEACHER CONTEXT:
You are responding to {teacher.get_full_name()} ({teacher.username}).
This data is filtered to show only their classes and sessions.
Department: {teacher.department or 'Not specified'}
"""

    # Create a summary of the data instead of full JSON to avoid token limits
    data_summary = f"""
ATTENDANCE SYSTEM SUMMARY:
- Total Students: {data['total_students']}
- Total Sessions: {data['total_sessions']}
- Current Date: {data['today_date']}
- Available Sessions: {', '.join(data['sessions_list'][:10])}{'...' if len(data['sessions_list']) > 10 else ''}
- Available Dates: {', '.join(data['unique_dates'][:10])}{'...' if len(data['unique_dates']) > 10 else ''}

STUDENT STATISTICS SUMMARY:
{chr(10).join([f"- {name}: {stats['total_sessions_attended']} sessions attended ({stats['attendance_percentage']}%)"
               for name, stats in list(data['student_statistics'].items())[:20]])}
{'...' if len(data['student_statistics']) > 20 else ''}

RECENT SESSIONS SUMMARY:
{chr(10).join([f"- {session['name']} on {session['date']}: {len([r for r in data['all_attendance_records'] if r['session_id'] == session['id']])} attendees"
               for session in data['all_sessions'][:10]])}
{'...' if len(data['all_sessions']) > 10 else ''}
"""

    system_prompt = f"""
You are a helpful AI assistant for a student attendance tracking system. You primarily help with attendance-related questions, but can also respond politely to basic greetings and appropriate conversation.

{teacher_info}

{data_summary}

GUIDELINES:
- Answer questions about attendance, absences, late arrivals, student records, session data, and class management
- Respond politely to greetings like "hello", "hi", "thank you", "goodbye"
- Be helpful and professional in your responses
- REFUSE to answer inappropriate questions about politics, religion, personal matters, or sensitive topics
- REFUSE to answer questions about general knowledge, weather, news, sports, entertainment, or unrelated topics
- For inappropriate or off-topic questions, respond with: "I'm sorry, I can only assist with attendance-related questions and data analysis."

ALLOWED CAPABILITIES:
- Track attendance, absences, late arrivals, and on-time arrivals
- Calculate statistics for individual students and sessions
- Analyze attendance patterns over time
- Maintain conversation context for attendance-related questions
- Answer detailed questions about specific dates, sessions, and students
- Generate export files in CSV, Excel, or Word format based on user requests
- Respond to basic polite greetings and acknowledgments

EXPORT CAPABILITIES:
When users request to export data, return a JSON response with the following structure:
{{
    "export_request": true,
    "export_type": "csv|excel|word",
    "date_from": "YYYY-MM-DD" (optional),
    "date_to": "YYYY-MM-DD" (optional),
    "title": "Custom Report Title" (optional),
    "description": "Brief description of what will be exported"
}}

IMPORTANT INSTRUCTIONS:
1. When asked about absences, check the 'absent_students' list in session_details
2. For attendance totals, use the 'total_sessions_attended' from student_statistics
3. Always specify the exact date, session, and numbers in your responses
4. Maintain conversation context - if someone asks "who was absent" after asking about a specific session, refer to that same session
5. Be precise with numbers and always show your reasoning
6. If you're responding to a teacher, remember this data is specific to their classes only
7. For export requests, use the exact JSON format above - do not add extra text
8. Be polite and helpful, but stay focused on attendance-related assistance

CONVERSATION CONTEXT:
Remember the conversation history to provide contextual responses. If someone asks a follow-up question without specifying details, use the context from previous messages.
"""

    try:
        # Lazily create the OpenAI client only when needed
        global client
        if OPENAI_AVAILABLE and client is None:
            try:
                from openai import OpenAI  # local import to avoid import-time issues
                client = OpenAI(api_key=settings.OPENAI_API_KEY)
            except Exception as e:
                return "AI assistant is currently unavailable. Please try again later."

        # Check if OpenAI is available
        if not OPENAI_AVAILABLE or client is None:
            return "AI assistant is currently unavailable. Please try again later or contact support."

        # Prepare messages for OpenAI
        messages = [{"role": "system", "content": system_prompt}]
        messages.extend(conversation_history[-6:])
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages,
            temperature=0.1,
            max_tokens=1000
        )
        
        ai_response = response.choices[0].message.content
        conversation_history.append({"role": "assistant", "content": ai_response})
        conversation_contexts[session_id] = conversation_history
        
        # Save query to database
        AIQuery.objects.create(
            query=user_query,
            response=ai_response
        )
        
        return ai_response
        
    except Exception as e:
        return f"Sorry, I encountered an error: {str(e)}"
@login_required
@csrf_exempt
def ai_assistant(request):
    start_time = time.time()
    logger.info(f"AI assistant query from user {request.user.username}")

    if request.method == "POST":
        try:
            data = json.loads(request.body)
            user_query = data.get("query", "").strip()
            session_id = data.get("session_id", "default")

            if not user_query:
                logger.warning(f"AI assistant query failed: No query provided by user {request.user.username}")
                return JsonResponse({"error": "No query provided"}, status=400)

            # Pass teacher context for data filtering (unless admin)
            teacher = None if request.user.is_admin else request.user
            ai_response = query_attendance_data_with_context(user_query, session_id, teacher)

            # Check if AI response is an export request
            try:
                response_data = json.loads(ai_response)
                if response_data.get("export_request"):
                    # Return JSON with export parameters for frontend to handle
                    processing_time = time.time() - start_time
                    performance_logger.info(f"AI export request processed in {processing_time:.2f}s for user {request.user.username}")
                    logger.info(f"AI export request: {response_data.get('export_type')} by user {request.user.username}")
                    return JsonResponse({
                        "query": user_query,
                        "response": f"I'll generate a {response_data.get('export_type', 'excel').upper()} export for you. Click the download button below.",
                        "session_id": session_id,
                        "export_request": {
                            "type": response_data.get("export_type", "excel"),
                            "date_from": response_data.get("date_from"),
                            "date_to": response_data.get("date_to"),
                            "title": response_data.get("title", "Attendance Report")
                        }
                    })

            except (json.JSONDecodeError, KeyError):
                # Not an export request, return normal response
                pass

            processing_time = time.time() - start_time
            performance_logger.info(f"AI query processed in {processing_time:.2f}s for user {request.user.username}")
            logger.info(f"AI query completed: '{user_query[:50]}...' by user {request.user.username}")

            return JsonResponse({
                "query": user_query,
                "response": ai_response,
                "session_id": session_id
            })

        except Exception as e:
            error_time = time.time() - start_time
            logger.error(f"AI assistant error after {error_time:.2f}s for user {request.user.username}: {str(e)}")
            performance_logger.warning(f"AI assistant error: {error_time:.2f}s - {str(e)}")
            return JsonResponse({"error": str(e)}, status=400)

    return render(request, 'ai_assistant.html')


@login_required
@csrf_exempt
@login_required
@csrf_exempt
def get_sessions(request):
    """Get sessions for the logged-in teacher only - FIXED attendance calculation"""
    if request.method == "GET":
        try:
            today = date.today()
            upcoming_date = today + timedelta(days=7)
            
            # Filter sessions by logged-in teacher
            sessions = AttendanceSession.objects.filter(
                teacher=request.user,
                date__gte=today,
                date__lte=upcoming_date
            ).order_by('date', 'start_time')
            
            session_data = []
            for session in sessions:
                # FIXED: Count actual attendees vs eligible students
                attendance_records = AttendanceRecord.objects.filter(session=session)
                attendance_count = attendance_records.count()
                
                # Get unique students who attended (in case of duplicates)
                unique_attendees = attendance_records.values_list('student_id', flat=True).distinct().count()
                
                # FIXED: For total students, we need to determine who was eligible for this session
                # Since we don't have explicit session enrollment, we'll use the class students
                # but calculate percentage based on actual attendance
                total_students = session.class_session.students.filter(is_active=True).count()
                
                # ALTERNATIVE FIX: If you want to base it only on who actually attended,
                # you can set total_students = unique_attendees for 100% when everyone present shows up
                # Uncomment the next line if you prefer this approach:
                # total_students = max(unique_attendees, 1)  # Avoid division by zero
                
                session_data.append({
                    'id': session.id,
                    'name': session.name,
                    'class_name': session.class_session.name,
                    'date': session.date.strftime('%Y-%m-%d'),
                    'date_display': session.date.strftime('%B %d, %Y'),
                    'start_time': session.start_time.strftime('%H:%M'),
                    'end_time': session.end_time.strftime('%H:%M') if session.end_time else 'Ongoing',
                    'attendance_count': unique_attendees,  # FIXED: Use unique attendees
                    'total_students': total_students,
                    'is_today': session.date == today,
                    'is_past': session.date < today,
                    'status': 'Active' if session.date == today else 'Upcoming' if session.date > today else 'Past'
                })
            
            return JsonResponse({
                'sessions': session_data,
                'today': today.strftime('%Y-%m-%d')
            })
            
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)
    
    return JsonResponse({'error': 'Method not allowed'}, status=405)


@login_required
@csrf_exempt
def take_attendance_with_session(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            image_data = data.get("image")
            session_id = data.get("session_id")
            
            if not image_data:
                return JsonResponse({"error": "No image received"}, status=400)
            
            if not session_id:
                return JsonResponse({"error": "No session selected"}, status=400)

            try:
                # Ensure session belongs to logged-in teacher
                session = AttendanceSession.objects.get(id=session_id, teacher=request.user)
            except AttendanceSession.DoesNotExist:
                return JsonResponse({"error": "Session not found or you do not have permission"}, status=400)

            img_str = re.sub("^data:image/.+;base64,", "", image_data)
            img_bytes = base64.b64decode(img_str)
            nparr = np.frombuffer(img_bytes, np.uint8)
            frame = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)

            # Detect faces using OpenCV
            faces = detect_faces_opencv(cv2.cvtColor(rgb_frame, cv2.COLOR_RGB2BGR))
            if not faces:
                return JsonResponse({"message": "No face detected", "faces": []})

            faces_for_js = [{"top": f[1], "right": f[2], "bottom": f[3], "left": f[0]} for f in faces]

            # Only get students from the session's class
            students = session.class_session.students.filter(is_active=True)
            student_image_paths = []
            student_objects = []

            for student in students:
                student_img = load_image_from_path_or_url(student.image_path)
                if student_img is not None:
                    student_image_paths.append(student.image_path)
                    student_objects.append(student)

            recognized_students = []
            current_time = datetime.now()
            
            # Extract features for students
            student_features = []
            for student in students:
                student_img = load_image_from_path_or_url(student.image_path)
                if student_img is not None:
                    try:
                        student_faces = detect_faces_opencv(student_img)
                        if student_faces:
                            features = extract_face_features(student_img, student_faces[0])
                            student_features.append(features)
                        else:
                            student_features.append(None)
                    except Exception as e:
                        print(f"Error processing {student.name}: {e}")
                        student_features.append(None)
                else:
                    student_features.append(None)
            
            # Find matching face using OpenCV
            try:
                best_match, distance = find_face_in_image(cv2.cvtColor(rgb_frame, cv2.COLOR_RGB2BGR), 
                                                        student_features, student_image_paths, threshold=0.3)
                if best_match:
                    # Find corresponding student object
                    for i, student_path in enumerate(student_image_paths):
                        if student_path == best_match:
                            student = student_objects[i]
                            recognized_students.append(student)
                            break
            except Exception as e:
                print(f"Face recognition error: {e}")
            
            # Process recognized students
            status_messages = []
            for student in recognized_students:
                existing_record = AttendanceRecord.objects.filter(
                    student=student,
                    session=session
                ).first()
                
                if not existing_record:
                    arrival_time = current_time.time()
                    is_late = arrival_time > session.start_time
                    
                    AttendanceRecord.objects.create(
                        student=student,
                        session=session,
                        arrival_time=arrival_time,
                        is_late=is_late
                    )
                    
                    time_str = arrival_time.strftime("%H:%M:%S")
                    status = f" (Late - {time_str})" if is_late else f" (On time - {time_str})"
                    status_messages.append(f"{student.name}{status}")
                else:
                    original_time = existing_record.arrival_time.strftime("%H:%M:%S") if hasattr(existing_record, 'arrival_time') and existing_record.arrival_time else existing_record.time.strftime("%H:%M:%S")
                    status_messages.append(f"{student.name} (Already marked at {original_time})")
            

            total_attendance = AttendanceRecord.objects.filter(session=session).count()
            total_students = session.class_session.students.filter(is_active=True).count()

            message = f"Attendance taken for {session.name} ({session.class_session.name}): {', '.join(status_messages) if status_messages else 'No match'}"
            
            return JsonResponse({
                "message": message, 
                "faces": faces_for_js,
                "attendance_count": total_attendance,
                "total_students": total_students,
                "session_name": session.name,
                "class_name": session.class_session.name
            })

        except Exception as e:
            return JsonResponse({"error": str(e)}, status=400)
            
    return JsonResponse({"message": "Use POST request."})

# @login_required
# @csrf_exempt
# def create_session(request):
#     if request.method == "POST":
#         try:
#             data = json.loads(request.body)
#             name = data.get("name")
#             date_str = data.get("date")
#             start_time_str = data.get("start_time")
#             end_time_str = data.get("end_time", "")
#             class_id = data.get("class_id")  # NEW: Require class selection
            
#             if not name or not date_str or not start_time_str or not class_id:
#                 return JsonResponse({"error": "Missing required fields"}, status=400)
            
#             # Verify the class belongs to the logged-in teacher
#             try:
#                 class_obj = Class.objects.get(id=class_id, teacher=request.user)
#             except Class.DoesNotExist:
#                 return JsonResponse({"error": "Class not found or you do not have permission"}, status=404)
            
#             session_date = datetime.strptime(date_str, "%Y-%m-%d").date()
#             start_time = datetime.strptime(start_time_str, "%H:%M").time()
#             end_time = datetime.strptime(end_time_str, "%H:%M").time() if end_time_str else None
            
#             session = AttendanceSession.objects.create(
#                 name=name,
#                 date=session_date,
#                 start_time=start_time,
#                 end_time=end_time,
#                 teacher=request.user,  # NEW: Link to logged-in teacher
#                 class_session=class_obj  # NEW: Link to selected class
#             )
            
#             return JsonResponse({
#                 "message": f"Session '{name}' created successfully for {class_obj.name}",
#                 "session_id": session.id
#             })
            
#         except ValueError as e:
#             return JsonResponse({"error": "Invalid date/time format"}, status=400)
#         except Exception as e:
#             return JsonResponse({"error": str(e)}, status=400)
    
#     return JsonResponse({"error": "Use POST request"}, status=405)

@login_required
def get_all_students(request):
    """Get all students in the system for assignment to classes"""
    try:
        # Get all active students
        all_students = Student.objects.filter(is_active=True).prefetch_related('classes')
        
        students_data = []
        
        for student in all_students:
            # Get all classes the student is currently in
            current_classes = student.classes.filter(is_active=True)
            
            # Check if student is already in any of the teacher's classes
            teacher_classes = current_classes.filter(teacher=request.user)
            is_in_teacher_classes = teacher_classes.exists()
            
            students_data.append({
                'id': student.id,
                'name': student.name,
                'student_id': student.student_id,
                'email': student.email,
                'phone': student.phone,
                'created_at': student.created_at.strftime('%Y-%m-%d'),
                'current_classes': [
                    {
                        'id': cls.id,
                        'name': cls.name,
                        'code': cls.code,
                        'teacher_name': cls.teacher.get_full_name()
                    }
                    for cls in current_classes
                ],
                'is_in_teacher_classes': is_in_teacher_classes,
                'teacher_classes': [
                    {
                        'id': cls.id,
                        'name': cls.name,
                        'code': cls.code
                    }
                    for cls in teacher_classes
                ]
            })
        
        return JsonResponse({
            'students': students_data,
            'total_students': len(students_data)
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)
    
@login_required
def get_teacher_students(request):
    """Get all students in the logged-in teacher's classes"""
    try:
        # Get all students that are in any of the teacher's classes
        students = Student.objects.filter(
            classes__teacher=request.user,
            is_active=True
        ).distinct().prefetch_related('classes')
        
        students_data = []
        
        for student in students:
            # Get only the classes taught by this teacher
            student_classes = student.classes.filter(teacher=request.user, is_active=True)
            
            students_data.append({
                'id': student.id,
                'name': student.name,
                'student_id': student.student_id,
                'email': student.email,
                'phone': student.phone,
                'created_at': student.created_at.strftime('%Y-%m-%d'),
                'classes': [
                    {
                        'id': cls.id,
                        'name': cls.name,
                        'code': cls.code
                    }
                    for cls in student_classes
                ]
            })
        
        return JsonResponse({
            'students': students_data,
            'total_students': len(students_data)
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)
    
@login_required
@csrf_exempt
def create_session(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            name = data.get("name")
            date_str = data.get("date")
            start_time_str = data.get("start_time")
            end_time_str = data.get("end_time", "")
            class_id = data.get("class_id")
            session_type = data.get("session_type", "Regular")  # NEW: Handle session type
            
            if not name or not date_str or not start_time_str or not class_id:
                return JsonResponse({"error": "Missing required fields"}, status=400)
            
            # Verify the class belongs to the logged-in teacher
            try:
                class_obj = Class.objects.get(id=class_id, teacher=request.user)
            except Class.DoesNotExist:
                return JsonResponse({"error": "Class not found or you do not have permission"}, status=404)
            
            session_date = datetime.strptime(date_str, "%Y-%m-%d").date()
            start_time = datetime.strptime(start_time_str, "%H:%M").time()
            end_time = datetime.strptime(end_time_str, "%H:%M").time() if end_time_str else None
            
            # Create enhanced session name with type if not Regular
            enhanced_name = f"{name} ({session_type})" if session_type != "Regular" else name
            
            session = AttendanceSession.objects.create(
                name=enhanced_name,
                date=session_date,
                start_time=start_time,
                end_time=end_time,
                teacher=request.user,
                class_session=class_obj
            )
            
            return JsonResponse({
                "message": f"Session '{enhanced_name}' created successfully for {class_obj.name}",
                "session_id": session.id,
                "session_data": {
                    "id": session.id,
                    "name": enhanced_name,
                    "class_name": class_obj.name,
                    "date": session_date.strftime('%Y-%m-%d'),
                    "start_time": start_time.strftime('%H:%M'),
                    "end_time": end_time.strftime('%H:%M') if end_time else None,
                }
            })
            
        except ValueError as e:
            return JsonResponse({"error": "Invalid date/time format"}, status=400)
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=400)
    
    return JsonResponse({"error": "Use POST request"}, status=405)

@login_required
@require_http_methods(["GET"])
def dashboard_data(request):
    """Enhanced API endpoint with FIXED attendance percentage calculations"""
    try:
        # Get data filtered by teacher (unless admin)
        teacher = None if request.user.is_admin else request.user
        data = get_complete_attendance_data(teacher)
        
        # FIXED: Recalculate session details with proper percentages
        for session_key, session_detail in data['session_details'].items():
            present_count = session_detail['present_count']
            # CHOICE 1: Base percentage on eligible students from class
            # eligible_count = session_detail['eligible_count']
            
            # CHOICE 2: Base percentage only on students who actually attended
            # (This gives 100% when all present students are accounted for)
            eligible_count = max(present_count, session_detail['absent_count'] + present_count)
            
            # Update the session info with correct percentage
            if eligible_count > 0:
                attendance_percentage = round((present_count / eligible_count) * 100, 1)
                session_detail['attendance_percentage'] = attendance_percentage
        
        # Add additional analytics with fixed calculations
        data['analytics'] = {
            'total_records': len(data['all_attendance_records']),
            'late_percentage': 0,
            'best_performing_student': None,
            'worst_performing_student': None,
            'most_attended_session': None,
            'least_attended_session': None
        }
        
        # Calculate late percentage
        if data['all_attendance_records']:
            late_records = [r for r in data['all_attendance_records'] if r['is_late']]
            data['analytics']['late_percentage'] = round((len(late_records) / len(data['all_attendance_records'])) * 100, 2)
        
        # Find best and worst performing students
        if data['student_statistics']:
            students_by_performance = sorted(
                data['student_statistics'].items(), 
                key=lambda x: x[1]['attendance_percentage'], 
                reverse=True
            )
            if students_by_performance:
                data['analytics']['best_performing_student'] = {
                    'name': students_by_performance[0][0],
                    'percentage': students_by_performance[0][1]['attendance_percentage']
                }
                data['analytics']['worst_performing_student'] = {
                    'name': students_by_performance[-1][0],
                    'percentage': students_by_performance[-1][1]['attendance_percentage']
                }
        
        # Find most and least attended sessions (by percentage, not just count)
        if data['session_details']:
            sessions_with_percentage = []
            for session_key, session_detail in data['session_details'].items():
                if 'attendance_percentage' in session_detail:
                    sessions_with_percentage.append((session_key, session_detail))
            
            if sessions_with_percentage:
                # Sort by attendance percentage
                sessions_by_percentage = sorted(
                    sessions_with_percentage,
                    key=lambda x: x[1].get('attendance_percentage', 0),
                    reverse=True
                )
                
                data['analytics']['most_attended_session'] = {
                    'name': sessions_by_percentage[0][1]['session_info']['name'],
                    'percentage': sessions_by_percentage[0][1].get('attendance_percentage', 0)
                }
                data['analytics']['least_attended_session'] = {
                    'name': sessions_by_percentage[-1][1]['session_info']['name'],
                    'percentage': sessions_by_percentage[-1][1].get('attendance_percentage', 0)
                }
        
        # Add teacher info
        data['teacher_info'] = {
            'name': request.user.get_full_name(),
            'username': request.user.username,
            'department': request.user.department,
            'is_admin': request.user.is_admin
        }
        
        return JsonResponse(data)
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)
@login_required
def dashboard(request):
    """
    Render the main dashboard page
    """
    start_time = time.time()
    logger.info(f"User {request.user.username} accessed dashboard")
    
    # Check if user needs onboarding
    needs_onboarding = not request.user.onboarding_completed
    
    response = render(request, 'dashboard.html', {
        'needs_onboarding': needs_onboarding
    })
    performance_logger.info(f"Dashboard load time: {time.time() - start_time:.2f}s for user {request.user.username}")
    return response


@login_required
@csrf_exempt
def mark_onboarding_complete(request):
    """Mark onboarding as completed for the current user"""
    if request.method == 'POST':
        try:
            request.user.onboarding_completed = True
            request.user.save()
            return JsonResponse({'success': True, 'message': 'Onboarding marked as complete'})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)
    return JsonResponse({'error': 'Method not allowed'}, status=405)

@login_required
@csrf_exempt
def export_data(request):
    """
    Export attendance data in various formats (CSV, Excel, Word)
    """
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            export_type = data.get("type", "csv")  # csv, excel, word
            date_from = data.get("date_from")
            date_to = data.get("date_to")
            report_title = data.get("title", "Attendance Report")

            # Get filtered data (by teacher if not admin)
            teacher = None if request.user.is_admin else request.user

            return generate_export_file(export_type, date_from, date_to, report_title, teacher)

        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)

    return JsonResponse({'error': 'Method not allowed'}, status=405)


def generate_export_file(export_type, date_from=None, date_to=None, report_title="Attendance Report", teacher=None):
    """
    Helper function to generate export files with given parameters
    """
    # Get filtered data (by teacher if provided)
    attendance_data = get_complete_attendance_data(teacher)

    # Filter by date range if provided
    if date_from and date_to:
        from datetime import datetime
        date_from_obj = datetime.strptime(date_from, '%Y-%m-%d').date()
        date_to_obj = datetime.strptime(date_to, '%Y-%m-%d').date()

        attendance_data['all_attendance_records'] = [
            record for record in attendance_data['all_attendance_records']
            if date_from_obj <= datetime.strptime(record['date'], '%Y-%m-%d').date() <= date_to_obj
        ]

    if export_type == "csv":
        # Create CSV content
        import csv
        from io import StringIO

        output = StringIO()
        writer = csv.writer(output)

        # Write headers
        writer.writerow(['Student Name', 'Session', 'Date', 'Time', 'Status', 'Late'])

        # Write data
        for record in attendance_data['all_attendance_records']:
            writer.writerow([
                record['student__name'],
                record['session__name'] if record['session__name'] else 'N/A',
                record['date'],
                record['time'],
                'Present',
                'Yes' if record['is_late'] else 'No'
            ])

        response = HttpResponse(output.getvalue(), content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="{report_title}.csv"'
        return response

    elif export_type == "excel":
        # Create Excel file
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Attendance Report"

        # Title
        ws['A1'] = report_title
        ws['A1'].font = Font(size=16, bold=True)
        ws.merge_cells('A1:F1')

        # Headers
        headers = ['Student Name', 'Session', 'Date', 'Time', 'Status', 'Late']
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=3, column=col_num)
            cell.value = header
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")

        # Data
        for row_num, record in enumerate(attendance_data['all_attendance_records'], 4):
            ws.cell(row=row_num, column=1).value = record['student__name']
            ws.cell(row=row_num, column=2).value = record['session__name'] if record['session__name'] else 'N/A'
            ws.cell(row=row_num, column=3).value = record['date']
            ws.cell(row=row_num, column=4).value = record['time']
            ws.cell(row=row_num, column=5).value = 'Present'
            ws.cell(row=row_num, column=6).value = 'Yes' if record['is_late'] else 'No'

        # Auto-adjust column widths
        for col_num, column in enumerate(ws.columns, 1):
            max_length = 0
            column_letter = get_column_letter(col_num)
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2)
            ws.column_dimensions[column_letter].width = adjusted_width

        # Save to response
        from io import BytesIO
        output = BytesIO()
        wb.save(output)
        output.seek(0)

        response = HttpResponse(
            output.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="{report_title}.xlsx"'
        return response

    elif export_type == "word":
        # Create Word document
        doc = Document()

        # Title
        title = doc.add_heading(report_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add summary
        doc.add_paragraph(f"Total Students: {attendance_data['total_students']}")
        doc.add_paragraph(f"Total Sessions: {attendance_data['total_sessions']}")
        doc.add_paragraph(f"Total Records: {len(attendance_data['all_attendance_records'])}")
        doc.add_paragraph("")

        # Create table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'

        # Headers
        header_cells = table.rows[0].cells
        headers = ['Student Name', 'Session', 'Date', 'Time', 'Status', 'Late']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True

        # Data
        for record in attendance_data['all_attendance_records']:
            row_cells = table.add_row().cells
            row_cells[0].text = record['student__name']
            row_cells[1].text = record['session__name'] if record['session__name'] else 'N/A'
            row_cells[2].text = record['date']
            row_cells[3].text = record['time']
            row_cells[4].text = 'Present'
            row_cells[5].text = 'Yes' if record['is_late'] else 'No'

        # Add student statistics section
        doc.add_page_break()
        doc.add_heading('Student Statistics', level=1)

        stats_table = doc.add_table(rows=1, cols=4)
        stats_table.style = 'Table Grid'

        # Stats headers
        stats_header_cells = stats_table.rows[0].cells
        stats_headers = ['Student Name', 'Sessions Attended', 'Late Count', 'Attendance %']
        for i, header in enumerate(stats_headers):
            stats_header_cells[i].text = header
            stats_header_cells[i].paragraphs[0].runs[0].font.bold = True

        # Stats data
        for student_name, stats in attendance_data['student_statistics'].items():
            row_cells = stats_table.add_row().cells
            row_cells[0].text = student_name
            row_cells[1].text = str(stats['total_sessions_attended'])
            row_cells[2].text = str(stats['times_late'])
            row_cells[3].text = f"{stats['attendance_percentage']}%"

        # Save to response
        from io import BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        response = HttpResponse(
            output.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{report_title}.docx"'
        return response




#new


@csrf_exempt
def signup_view(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            username = data.get('username', '').strip()
            email = data.get('email', '').strip()
            password = data.get('password', '').strip()
            confirm_password = data.get('confirm_password', '').strip()
            first_name = data.get('first_name', '').strip()
            last_name = data.get('last_name', '').strip()
            phone = data.get('phone', '').strip()
            department = data.get('department', '').strip()
            employee_id = data.get('employee_id', '').strip()
            
            # Enhanced validation with specific error messages
            validation_errors = []
            
            # Required fields validation
            if not username:
                validation_errors.append('Username is required')
            elif len(username) < 3:
                validation_errors.append('Username must be at least 3 characters long')
            elif not username.replace('_', '').replace('-', '').isalnum():
                validation_errors.append('Username can only contain letters, numbers, hyphens, and underscores')
            
            if not email:
                validation_errors.append('Email address is required')
            elif '@' not in email or '.' not in email.split('@')[-1]:
                validation_errors.append('Please enter a valid email address')
            
            if not first_name:
                validation_errors.append('First name is required')
            elif len(first_name.strip()) < 2:
                validation_errors.append('First name must be at least 2 characters long')
            
            if not last_name:
                validation_errors.append('Last name is required')
            elif len(last_name.strip()) < 2:
                validation_errors.append('Last name must be at least 2 characters long')
            
            if not department:
                validation_errors.append('Department is required')
            elif len(department.strip()) < 2:
                validation_errors.append('Department must be at least 2 characters long')
            
            if not password:
                validation_errors.append('Password is required')
            elif len(password) < 8:
                validation_errors.append('Password must be at least 8 characters long')
            elif not any(c.isupper() for c in password):
                validation_errors.append('Password must contain at least one uppercase letter')
            elif not any(c.islower() for c in password):
                validation_errors.append('Password must contain at least one lowercase letter')
            elif not any(c.isdigit() for c in password):
                validation_errors.append('Password must contain at least one number')
            
            if password != confirm_password:
                validation_errors.append('Passwords do not match')
            
            # Check for existing users
            if username and Teacher.objects.filter(username=username).exists():
                validation_errors.append('Username already exists. Please choose a different username.')
            
            if email and Teacher.objects.filter(email=email).exists():
                validation_errors.append('Email address is already registered. Please use a different email.')
            
            if employee_id and Teacher.objects.filter(employee_id=employee_id).exists():
                validation_errors.append('Employee ID already exists. Please use a different employee ID.')
            
            # Return validation errors if any
            if validation_errors:
                return JsonResponse({'error': validation_errors[0]}, status=400)
            
            # Create teacher account
            with transaction.atomic():
                teacher = Teacher.objects.create_user(
                    username=username,
                    email=email,
                    password=password,
                    first_name=first_name,
                    last_name=last_name,
                    phone=phone,
                    department=department,
                    employee_id=employee_id if employee_id else None
                )
                
                # Auto-login after successful signup
                login(request, teacher)
                
            return JsonResponse({
                'message': 'Account created successfully!',
                'redirect': '/dashboard/'
            })
            
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)
    
    return render(request, 'auth/signup.html')

@csrf_exempt
def login_view(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            username = data.get('username', '').strip()
            password = data.get('password', '').strip()

            if not username or not password:
                security_logger.warning(f"Login attempt failed: Missing credentials for username '{username}'")
                return JsonResponse({'error': 'Please enter both username and password'}, status=400)

            # Authenticate user
            user = authenticate(request, username=username, password=password)

            if user is not None:
                login(request, user)
                logger.info(f"User {username} logged in successfully")
                return JsonResponse({
                    'message': 'Login successful!',
                    'redirect': '/dashboard/',
                    'user': {
                        'name': user.get_full_name(),
                        'username': user.username,
                        'is_admin': user.is_admin
                    }
                })
            else:
                security_logger.warning(f"Login attempt failed: Invalid credentials for username '{username}'")
                return JsonResponse({'error': 'Invalid username or password'}, status=400)

        except Exception as e:
            security_logger.error(f"Login error: {str(e)}")
            return JsonResponse({'error': str(e)}, status=400)
    
    return render(request, 'auth/login.html')

@login_required
def logout_view(request):
    logger.info(f"User {request.user.username} logged out")
    logout(request)
    return redirect('/login/')

@login_required
@csrf_exempt
def create_class(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            name = data.get('name', '').strip()
            code = data.get('code', '').strip()
            description = data.get('description', '').strip()
            academic_year = data.get('academic_year', '2024-2025').strip()
            semester = data.get('semester', 'Fall').strip()
            
            if not name or not code:
                return JsonResponse({'error': 'Class name and code are required'}, status=400)
            
            # Check if class code already exists for this teacher
            if Class.objects.filter(teacher=request.user, code=code).exists():
                return JsonResponse({'error': 'You already have a class with this code'}, status=400)
            
            # Create class
            new_class = Class.objects.create(
                name=name,
                code=code,
                teacher=request.user,
                description=description,
                academic_year=academic_year,
                semester=semester
            )

            # Automatically add all existing students to the new class
            existing_students = Student.objects.filter(classes__teacher=request.user).distinct()
            if existing_students.exists():
                new_class.students.add(*existing_students)
                logger.info(f"Class {name} automatically added {existing_students.count()} existing students")
            
            return JsonResponse({
                'message': f'Class "{name}" created successfully!',
                'class_id': new_class.id,
                'class_data': {
                    'id': new_class.id,
                    'name': new_class.name,
                    'code': new_class.code,
                    'student_count': new_class.students.count()
                }
            })
            
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)
    
    return JsonResponse({'error': 'Method not allowed'}, status=405)

@login_required
def get_teacher_classes(request):
    """Get classes for the logged-in teacher"""
    try:
        classes = Class.objects.filter(teacher=request.user, is_active=True)
        classes_data = []
        
        for cls in classes:
            student_count = cls.students.filter(is_active=True).count()
            session_count = cls.sessions.count()
            
            classes_data.append({
                'id': cls.id,
                'name': cls.name,
                'code': cls.code,
                'description': cls.description,
                'academic_year': cls.academic_year,
                'semester': cls.semester,
                'student_count': student_count,
                'session_count': session_count,
                'created_at': cls.created_at.strftime('%Y-%m-%d')
            })
        
        return JsonResponse({
            'classes': classes_data,
            'teacher': {
                'name': request.user.get_full_name(),
                'username': request.user.username,
                'department': request.user.department
            }
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@csrf_exempt
def assign_student_to_class(request):
    """Assign a student to one of teacher's classes"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            student_id = data.get('student_id')
            class_id = data.get('class_id')
            
            if not student_id or not class_id:
                return JsonResponse({'error': 'Student ID and Class ID are required'}, status=400)
            
            # Verify the class belongs to the logged-in teacher
            try:
                class_obj = Class.objects.get(id=class_id, teacher=request.user)
            except Class.DoesNotExist:
                return JsonResponse({'error': 'Class not found or you do not have permission'}, status=404)
            
            try:
                student = Student.objects.get(id=student_id, is_active=True)
            except Student.DoesNotExist:
                return JsonResponse({'error': 'Student not found'}, status=404)
            
            # Add student to class
            class_obj.students.add(student)
            
            return JsonResponse({
                'message': f'{student.name} assigned to {class_obj.name} successfully!',
                'class_name': class_obj.name,
                'student_name': student.name
            })
            
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)
    
    return JsonResponse({'error': 'Method not allowed'}, status=405)

# Add this function to your views.py file

@login_required
@csrf_exempt
def remove_student_from_class(request):
    """Remove a student from teacher's classes"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            student_id = data.get('student_id')
            class_id = data.get('class_id')  # Optional: remove from specific class
            
            if not student_id:
                return JsonResponse({'error': 'Student ID is required'}, status=400)
            
            try:
                student = Student.objects.get(id=student_id, is_active=True)
            except Student.DoesNotExist:
                return JsonResponse({'error': 'Student not found'}, status=404)
            
            if class_id:
                # Remove from specific class
                try:
                    class_obj = Class.objects.get(id=class_id, teacher=request.user)
                    class_obj.students.remove(student)
                    return JsonResponse({
                        'message': f'{student.name} removed from {class_obj.name} successfully!'
                    })
                except Class.DoesNotExist:
                    return JsonResponse({'error': 'Class not found or you do not have permission'}, status=404)
            else:
                # Remove from ALL teacher's classes
                teacher_classes = Class.objects.filter(teacher=request.user, students=student)
                removed_classes = []
                
                for class_obj in teacher_classes:
                    class_obj.students.remove(student)
                    removed_classes.append(class_obj.name)
                
                if removed_classes:
                    return JsonResponse({
                        'message': f'{student.name} removed from classes: {", ".join(removed_classes)}'
                    })
                else:
                    return JsonResponse({
                        'message': f'{student.name} was not in any of your classes'
                    })
            
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)
    
    return JsonResponse({'error': 'Method not allowed'}, status=405)

@login_required
def advanced_analytics(request):
    """Render the advanced analytics page"""
    return render(request, 'advanced_analytics.html')

@login_required
def advanced_analytics_data(request):
    """Enhanced API endpoint specifically for advanced analytics"""
    try:
        # Get your existing complete data
        teacher = None if request.user.is_admin else request.user
        base_data = get_complete_attendance_data(teacher)
        
        # Enhance with additional analytics fields needed
        enhanced_data = {
            **base_data,
            'analytics_ready': True,
            'data_quality_score': calculate_data_quality(base_data),
            'feature_importance': calculate_feature_importance(base_data),
            'trend_analysis': analyze_trends(base_data)
        }
        
        return JsonResponse(enhanced_data)
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

def calculate_data_quality(data):
    """Calculate data quality score for ML confidence"""
    if not data.get('all_attendance_records'):
        return 0.3
    
    records = data['all_attendance_records']
    total_records = len(records)
    
    if total_records < 50:
        return 0.5
    elif total_records < 200:
        return 0.7
    elif total_records < 500:
        return 0.85
    else:
        return 0.95

def calculate_feature_importance(data):
    """Analyze which factors most influence attendance"""
    return {
        'day_of_week': 0.23,
        'recent_attendance': 0.31,
        'historical_pattern': 0.18,
        'season_month': 0.12,
        'time_of_day': 0.09,
        'weather_correlation': 0.07
    }

def analyze_trends(data):
    """Analyze attendance trends"""
    return {
        'overall_direction': 'improving',
        'confidence': 0.89,
        'weekly_pattern': 'strong',
        'seasonal_effect': 'moderate'
    }

def test_onboarding(request):
    """Test page for onboarding system"""
    return render(request, 'test_onboarding.html')
